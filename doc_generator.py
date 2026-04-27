"""
DOCX document generator module.

This module handles filling DOCX templates with parsed data from RedCheck reports.
It uses python-docx library to manipulate Word documents.
"""

from typing import Dict, List, Any, Optional
from pathlib import Path
import logging
from datetime import datetime

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from parsers import ParsedData, Host, PortScan, Vulnerability

logger = logging.getLogger(__name__)


class DocumentGenerator:
    """Generates security protocol documents from parsed report data."""

    def __init__(self, template_path: str):
        """
        Initialize document generator with template path.

        Args:
            template_path: Path to the DOCX template file
        """
        self.template_path = template_path
        self.document: Optional[Document] = None
        self.errors: List[str] = []
        self.warnings: List[str] = []

    def load_template(self) -> bool:
        """
        Load the DOCX template.

        Returns:
            True if successful, False otherwise
        """
        try:
            if not Path(self.template_path).exists():
                self._log_error(f"Template file not found: {self.template_path}")
                return False
            
            self.document = Document(self.template_path)
            logger.info(f"Template loaded successfully: {self.template_path}")
            return True
        except Exception as e:
            self._log_error(f"Error loading template: {e}")
            return False

    def fill_document(self, data: ParsedData, output_path: str) -> bool:
        """
        Fill the document with parsed data and save it.

        Args:
            data: ParsedData object containing all report information
            output_path: Path for the output DOCX file

        Returns:
            True if successful, False otherwise
        """
        if not self.document:
            if not self.load_template():
                return False

        try:
            # Fill different sections based on table placeholders or bookmarks
            self._fill_inventory_tables(data.hosts)
            self._fill_pentest_tables(data.port_scans)
            self._fill_vulnerability_tables(data.vulnerabilities)
            self._fill_metadata_section(data.scan_metadata)

            # Save the document
            self.document.save(output_path)
            logger.info(f"Document saved successfully: {output_path}")
            return True

        except Exception as e:
            self._log_error(f"Error filling document: {e}")
            return False

    def _find_table_by_header(self, header_text: str) -> Optional[Table]:
        """
        Find a table by its header text.

        Args:
            header_text: Text to search for in table headers

        Returns:
            Table object if found, None otherwise
        """
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if header_text.lower() in cell.text.lower():
                        return table
        return None

    def _find_paragraph_by_text(self, search_text: str) -> Optional[Paragraph]:
        """
        Find a paragraph containing specific text.

        Args:
            search_text: Text to search for

        Returns:
            Paragraph object if found, None otherwise
        """
        for para in self.document.paragraphs:
            if search_text.lower() in para.text.lower():
                return para
        return None

    def _fill_inventory_tables(self, hosts: List[Host]):
        """Fill inventory-related tables with host data."""
        if not hosts:
            logger.info("No inventory data to fill")
            return

        # Try to find tables by common header names
        host_table = self._find_table_by_header("хост") or \
                    self._find_table_by_header("host") or \
                    self._find_table_by_header("оборудование") or \
                    self._find_table_by_header("inventory")

        if host_table:
            self._populate_host_table(host_table, hosts)
        else:
            self._log_warning("No matching table found for inventory data")

        # Fill software table if exists
        software_table = self._find_table_by_header("программное обеспечение") or \
                        self._find_table_by_header("software") or \
                        self._find_table_by_header("по")

        if software_table:
            self._populate_software_table(software_table, hosts)

        # Fill users table if exists
        users_table = self._find_table_by_header("пользователь") or \
                     self._find_table_by_header("user") or \
                     self._find_table_by_header("учетная запись")

        if users_table:
            self._populate_users_table(users_table, hosts)

        # Fill services table if exists
        services_table = self._find_table_by_header("служба") or \
                        self._find_table_by_header("service") or \
                        self._find_table_by_header("процесс")

        if services_table:
            self._populate_services_table(services_table, hosts)

    def _populate_host_table(self, table: Table, hosts: List[Host]):
        """Populate host inventory table."""
        # Clear existing rows except header
        self._clear_table_rows(table, keep_header=True)

        for host in hosts:
            row = table.add_row()
            cells = row.cells

            # Map host data to table columns based on header text
            headers = [cell.text.lower() for cell in table.rows[0].cells]

            for i, header in enumerate(headers):
                if i < len(cells):
                    value = self._get_host_field(host, header)
                    cells[i].text = str(value) if value else ""

    def _get_host_field(self, host: Host, header: str) -> Any:
        """Get host field value based on header text."""
        header_lower = header.lower()

        if 'ip' in header_lower or 'адрес' in header_lower:
            return host.ip
        elif 'hostname' in header_lower or 'имя' in header_lower and 'хост' in header_lower:
            return host.hostname
        elif 'os' in header_lower or 'операционная система' in header_lower:
            return host.os
        elif 'arch' in header_lower or 'архитектура' in header_lower:
            return host.architecture
        elif 'domain' in header_lower or 'домен' in header_lower:
            return host.domain_role
        elif 'mac' in header_lower:
            return host.mac_address
        elif 'network' in header_lower or 'интерфейс' in header_lower:
            return self._format_interfaces(host.network_interfaces)
        elif 'user' in header_lower:
            return self._format_list([u.get('name', '') for u in host.users])
        elif 'service' in header_lower:
            return self._format_list([s.get('name', '') for s in host.services])
        elif 'software' in header_lower or 'по' in header_lower:
            return self._format_software_list(host.installed_software)
        elif 'update' in header_lower or 'kb' in header_lower:
            return self._format_updates_list(host.updates)

        return ""

    def _populate_software_table(self, table: Table, hosts: List[Host]):
        """Populate software installation table."""
        self._clear_table_rows(table, keep_header=True)

        headers = [cell.text.lower() for cell in table.rows[0].cells]

        for host in hosts:
            for software in host.installed_software:
                row = table.add_row()
                cells = row.cells

                for i, header in enumerate(headers):
                    if i < len(cells):
                        value = self._get_software_field(software, host, header)
                        cells[i].text = str(value) if value else ""

    def _get_software_field(self, software: Dict[str, Any], host: Host, header: str) -> Any:
        """Get software field value based on header text."""
        header_lower = header.lower()

        if 'name' in header_lower or 'название' in header_lower or 'по' in header_lower:
            return software.get('name', '')
        elif 'version' in header_lower or 'версия' in header_lower:
            return software.get('version', '')
        elif 'date' in header_lower or 'дата' in header_lower:
            return software.get('date', '')
        elif 'vendor' in header_lower or 'производитель' in header_lower:
            return software.get('vendor', '')
        elif 'host' in header_lower or 'ip' in header_lower:
            return host.ip
        elif 'hostname' in header_lower:
            return host.hostname

        return ""

    def _populate_users_table(self, table: Table, hosts: List[Host]):
        """Populate users table."""
        self._clear_table_rows(table, keep_header=True)

        headers = [cell.text.lower() for cell in table.rows[0].cells]

        for host in hosts:
            for user in host.users:
                row = table.add_row()
                cells = row.cells

                for i, header in enumerate(headers):
                    if i < len(cells):
                        value = self._get_user_field(user, host, header)
                        cells[i].text = str(value) if value else ""

    def _get_user_field(self, user: Dict[str, Any], host: Host, header: str) -> Any:
        """Get user field value based on header text."""
        header_lower = header.lower()

        if 'name' in header_lower or 'username' in header_lower or 'пользователь' in header_lower:
            return user.get('name', '')
        elif 'type' in header_lower or 'тип' in header_lower:
            return user.get('type', '')
        elif 'enabled' in header_lower or 'active' in header_lower or 'статус' in header_lower:
            return user.get('enabled', '')
        elif 'last_login' in header_lower or 'last logon' in header_lower or 'вход' in header_lower:
            return user.get('last_login', '')
        elif 'host' in header_lower or 'ip' in header_lower:
            return host.ip
        elif 'hostname' in header_lower:
            return host.hostname

        return ""

    def _populate_services_table(self, table: Table, hosts: List[Host]):
        """Populate services table."""
        self._clear_table_rows(table, keep_header=True)

        headers = [cell.text.lower() for cell in table.rows[0].cells]

        for host in hosts:
            for service in host.services:
                row = table.add_row()
                cells = row.cells

                for i, header in enumerate(headers):
                    if i < len(cells):
                        value = self._get_service_field(service, host, header)
                        cells[i].text = str(value) if value else ""

    def _get_service_field(self, service: Dict[str, Any], host: Host, header: str) -> Any:
        """Get service field value based on header text."""
        header_lower = header.lower()

        if 'name' in header_lower or 'название' in header_lower or 'служба' in header_lower:
            return service.get('name', '')
        elif 'status' in header_lower or 'статус' in header_lower:
            return service.get('status', '')
        elif 'startup' in header_lower or 'запуск' in header_lower:
            return service.get('startup', '')
        elif 'path' in header_lower or 'путь' in header_lower:
            return service.get('path', '')
        elif 'host' in header_lower or 'ip' in header_lower:
            return host.ip
        elif 'hostname' in header_lower:
            return host.hostname

        return ""

    def _fill_pentest_tables(self, port_scans: List[PortScan]):
        """Fill pentest-related tables with port scan data."""
        if not port_scans:
            logger.info("No pentest data to fill")
            return

        # Find ports table
        ports_table = self._find_table_by_header("порт") or \
                     self._find_table_by_header("port") or \
                     self._find_table_by_header("open port") or \
                     self._find_table_by_header("открытый порт")

        if ports_table:
            self._populate_ports_table(ports_table, port_scans)

        # Find SMB table
        smb_table = self._find_table_by_header("smb") or \
                   self._find_table_by_header("SMB")

        if smb_table:
            self._populate_smb_table(smb_table, port_scans)

    def _populate_ports_table(self, table: Table, port_scans: List[PortScan]):
        """Populate open ports table."""
        self._clear_table_rows(table, keep_header=True)

        headers = [cell.text.lower() for cell in table.rows[0].cells]

        for scan in port_scans:
            for port in scan.open_ports:
                row = table.add_row()
                cells = row.cells

                for i, header in enumerate(headers):
                    if i < len(cells):
                        value = self._get_port_field(port, scan, header)
                        cells[i].text = str(value) if value else ""

    def _get_port_field(self, port: Dict[str, Any], scan: PortScan, header: str) -> Any:
        """Get port field value based on header text."""
        header_lower = header.lower()

        if 'port' in header_lower or 'порт' in header_lower:
            return port.get('port', '')
        elif 'protocol' in header_lower or 'протокол' in header_lower:
            return port.get('protocol', '')
        elif 'state' in header_lower or 'status' in header_lower or 'состояние' in header_lower:
            return port.get('state', '')
        elif 'service' in header_lower or 'сервис' in header_lower:
            return self._find_service_for_port(scan.services, port.get('port'))
        elif 'host' in header_lower or 'ip' in header_lower:
            return scan.ip
        elif 'hostname' in header_lower:
            return scan.hostname

        return ""

    def _find_service_for_port(self, services: List[Dict[str, Any]], port: str) -> str:
        """Find service name for a given port."""
        for svc in services:
            if str(svc.get('port', '')) == str(port):
                return svc.get('name', '')
        return ""

    def _populate_smb_table(self, table: Table, port_scans: List[PortScan]):
        """Populate SMB configuration table."""
        self._clear_table_rows(table, keep_header=True)

        headers = [cell.text.lower() for cell in table.rows[0].cells]

        for scan in port_scans:
            if scan.smb_info:
                row = table.add_row()
                cells = row.cells

                for i, header in enumerate(headers):
                    if i < len(cells):
                        value = self._get_smb_field(scan.smb_info, scan, header)
                        cells[i].text = str(value) if value else ""

    def _get_smb_field(self, smb_info: Dict[str, Any], scan: PortScan, header: str) -> Any:
        """Get SMB field value based on header text."""
        header_lower = header.lower()

        if 'signature' in header_lower and ('required' in header_lower or 'требуется' in header_lower):
            return smb_info.get('signature_required', '')
        elif 'signature' in header_lower and ('enabled' in header_lower or 'включено' in header_lower):
            return smb_info.get('signature_enabled', '')
        elif 'domain' in header_lower or 'домен' in header_lower:
            return smb_info.get('domain', '')
        elif 'workgroup' in header_lower:
            return smb_info.get('workgroup', '')
        elif 'os' in header_lower or 'версия' in header_lower:
            return smb_info.get('os_version', '')
        elif 'server' in header_lower or 'тип сервера' in header_lower:
            return smb_info.get('server_type', '')
        elif 'lanman' in header_lower:
            return smb_info.get('lanman_version', '')
        elif 'ntlm' in header_lower:
            return smb_info.get('ntlm_version', '')
        elif 'host' in header_lower or 'ip' in header_lower:
            return scan.ip
        elif 'hostname' in header_lower:
            return scan.hostname

        return ""

    def _fill_vulnerability_tables(self, vulnerabilities: List[Vulnerability]):
        """Fill vulnerability-related tables."""
        if not vulnerabilities:
            logger.info("No vulnerability data to fill")
            return

        vuln_table = self._find_table_by_header("уязвим") or \
                    self._find_table_by_header("vulnerability") or \
                    self._find_table_by_header("vuln") or \
                    self._find_table_by_header("cvss")

        if vuln_table:
            self._populate_vulnerability_table(vuln_table, vulnerabilities)

    def _populate_vulnerability_table(self, table: Table, vulnerabilities: List[Vulnerability]):
        """Populate vulnerability table."""
        self._clear_table_rows(table, keep_header=True)

        headers = [cell.text.lower() for cell in table.rows[0].cells]

        for vuln in vulnerabilities:
            row = table.add_row()
            cells = row.cells

            for i, header in enumerate(headers):
                if i < len(cells):
                    value = self._get_vulnerability_field(vuln, header)
                    cells[i].text = str(value) if value else ""

    def _get_vulnerability_field(self, vuln: Vulnerability, header: str) -> Any:
        """Get vulnerability field value based on header text."""
        header_lower = header.lower()

        if 'id' in header_lower:
            return vuln.id
        elif 'cve' in header_lower:
            return vuln.cve_id
        elif 'name' in header_lower or 'название' in header_lower or 'уязвимость' in header_lower:
            return vuln.name
        elif 'cvss' in header_lower:
            return vuln.cvss_score
        elif 'severity' in header_lower or 'critical' in header_lower or 'критичность' in header_lower or 'важность' in header_lower:
            return vuln.severity
        elif 'exploit' in header_lower and ('available' in header_lower or 'доступен' in header_lower):
            return 'Yes' if vuln.exploit_available else 'No'
        elif 'exploitation' in header_lower or 'exploited' in header_lower or 'эксплуатация' in header_lower:
            return vuln.exploitation_status
        elif 'description' in header_lower or 'описание' in header_lower:
            return vuln.description
        elif 'remediation' in header_lower or 'solution' in header_lower or 'решение' in header_lower or 'устранение' in header_lower:
            return vuln.remediation
        elif 'target' in header_lower or 'ip' in header_lower:
            return vuln.target_ip
        elif 'port' in header_lower:
            return vuln.target_port

        return ""

    def _fill_metadata_section(self, metadata: Dict[str, Any]):
        """Fill metadata section in the document."""
        if not metadata:
            return

        # Look for placeholder text patterns like {{scan_type}}, ${profile}, etc.
        for para in self.document.paragraphs:
            text = para.text
            updated = False

            for key, value in metadata.items():
                placeholders = [
                    f"{{{{{key}}}}}",
                    f"${{{key}}}",
                    f"[{key}]",
                    f"<{key}>",
                ]
                for placeholder in placeholders:
                    if placeholder in text:
                        text = text.replace(placeholder, str(value))
                        updated = True

            if updated:
                # Update paragraph text while preserving formatting
                if para.runs:
                    para.runs[0].text = text
                else:
                    para.text = text

    def _clear_table_rows(self, table: Table, keep_header: bool = True):
        """
        Clear all rows from a table, optionally keeping the header.

        Args:
            table: Table to clear
            keep_header: If True, keep the first row (header)
        """
        start_index = 1 if keep_header else 0
        for i in range(len(table.rows) - 1, start_index - 1, -1):
            table._tbl.remove(table.rows[i]._tr)

    def _format_list(self, items: List[str], separator: str = ", ") -> str:
        """Format a list of items into a string."""
        return separator.join(filter(None, items))

    def _format_interfaces(self, interfaces: List[Dict[str, Any]]) -> str:
        """Format network interfaces into a readable string."""
        if not interfaces:
            return ""
        
        formatted = []
        for iface in interfaces:
            parts = []
            if iface.get('ip'):
                parts.append(f"{iface['ip']}/{iface.get('mask', '')}")
            if iface.get('name'):
                parts.insert(0, f"{iface['name']}:")
            formatted.append(" ".join(parts))
        
        return "; ".join(formatted)

    def _format_software_list(self, software: List[Dict[str, Any]]) -> str:
        """Format software list into a readable string."""
        if not software:
            return ""
        
        formatted = []
        for sw in software:
            name = sw.get('name', '')
            version = sw.get('version', '')
            if name:
                entry = f"{name} ({version})" if version else name
                formatted.append(entry)
        
        return "; ".join(formatted[:10])  # Limit to first 10 items

    def _format_updates_list(self, updates: List[Dict[str, Any]]) -> str:
        """Format updates list into a readable string."""
        if not updates:
            return ""
        
        formatted = []
        for upd in updates:
            kb_id = upd.get('id', '')
            name = upd.get('name', '')
            if kb_id:
                entry = f"{kb_id}"
                if name:
                    entry += f": {name}"
                formatted.append(entry)
        
        return "; ".join(formatted[:20])  # Limit to first 20 items

    def _log_error(self, message: str):
        """Log an error message."""
        logger.error(message)
        self.errors.append(message)

    def _log_warning(self, message: str):
        """Log a warning message."""
        logger.warning(message)
        self.warnings.append(message)

    def get_errors(self) -> List[str]:
        """Get list of errors."""
        return self.errors.copy()

    def get_warnings(self) -> List[str]:
        """Get list of warnings."""
        return self.warnings.copy()
