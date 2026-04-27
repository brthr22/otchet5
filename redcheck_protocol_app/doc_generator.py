"""
Document generator for RedCheck Protocol Generator.
Fills DOCX template tables with parsed data.
"""

import logging
from pathlib import Path
from typing import List, Dict, Any, Optional
from copy import deepcopy

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from core.models import Host, PortScan, Vulnerability, ScanMetadata


logger = logging.getLogger(__name__)


class DocumentGenerator:
    """Generates protocol document from template and parsed data."""
    
    # Table type identifiers (search patterns in table captions or first cell)
    TABLE_PATTERNS = {
        'hosts': ['инвентаризация', 'хосты', 'hosts', 'inventory', 'сведения об ос'],
        'ports': ['порты', 'ports', 'открытые порты', 'port scan', 'сетевые сервисы'],
        'vulnerabilities': ['уязвимости', 'vulnerabilities', 'vulns', 'cvss', 'cve'],
        'software': ['по', 'software', 'программное обеспечение', 'приложения'],
        'services': ['службы', 'services', 'сервисы', 'процессы'],
        'users': ['пользователи', 'users', 'учетные записи', 'accounts'],
        'network': ['сеть', 'network', 'интерфейсы', 'interfaces', 'ip'],
        'updates': ['обновления', 'updates', 'kb', 'патчи', 'patches'],
    }
    
    def __init__(self, template_path: Path):
        self.template_path = template_path
        self.doc = None
    
    def generate(self, output_path: Path, hosts_data: List[Dict[str, Any]]):
        """
        Generate filled document from template.
        
        Args:
            output_path: Path to save the generated document
            hosts_data: List of dicts with keys: host, port_scan, vulnerabilities, scan_metadata, errors
        """
        if not self.template_path.exists():
            raise FileNotFoundError(f"Template not found: {self.template_path}")
        
        # Load template
        self.doc = Document(str(self.template_path))
        logger.info(f"Loaded template: {self.template_path}")
        
        # Process each host
        for i, data in enumerate(hosts_data):
            logger.info(f"Processing host {i + 1}/{len(hosts_data)}")
            self._process_host(data, i)
        
        # Save document
        self.doc.save(str(output_path))
        logger.info(f"Document saved: {output_path}")
    
    def _process_host(self, data: Dict[str, Any], host_index: int):
        """Process a single host's data and fill tables."""
        host = data.get('host')
        port_scan = data.get('port_scan')
        vulnerabilities = data.get('vulnerabilities', [])
        scan_metadata = data.get('scan_metadata')
        errors = data.get('errors', [])
        
        # Find and fill tables
        tables_found = {key: False for key in self.TABLE_PATTERNS}
        
        for table in self.doc.tables:
            table_type = self._identify_table(table)
            
            if table_type is None:
                continue
            
            if table_type == 'hosts' and host:
                self._fill_host_table(table, host, host_index)
                tables_found['hosts'] = True
                
            elif table_type == 'ports' and port_scan:
                self._fill_ports_table(table, port_scan, host_index)
                tables_found['ports'] = True
                
            elif table_type == 'vulnerabilities' and vulnerabilities:
                self._fill_vulns_table(table, vulnerabilities, host_index)
                tables_found['vulnerabilities'] = True
                
            elif table_type == 'software' and host and host.installed_software:
                self._fill_software_table(table, host.installed_software, host_index)
                tables_found['software'] = True
                
            elif table_type == 'services' and host and host.services:
                self._fill_services_table(table, host.services, host_index)
                tables_found['services'] = True
                
            elif table_type == 'users' and host and host.users_groups:
                self._fill_users_table(table, host.users_groups, host_index)
                tables_found['users'] = True
                
            elif table_type == 'network' and host and host.interfaces:
                self._fill_network_table(table, host.interfaces, host_index)
                tables_found['network'] = True
                
            elif table_type == 'updates' and host and host.updates_kb:
                self._fill_updates_table(table, host.updates_kb, host_index)
                tables_found['updates'] = True
        
        # Log warnings for missing data
        for table_type, found in tables_found.items():
            if not found:
                logger.debug(f"Table type '{table_type}' not filled (may be missing data or table)")
        
        # Log errors if any
        if errors:
            for error in errors:
                logger.warning(f"Host {host_index + 1}: {error}")
    
    def _identify_table(self, table) -> Optional[str]:
        """Identify table type based on content."""
        # Check all cells in header row and first few rows
        try:
            header_text = ""
            for row_idx in range(min(2, len(table.rows))):
                for cell in table.rows[row_idx].cells:
                    header_text += " " + cell.text.lower()
            
            for table_type, patterns in self.TABLE_PATTERNS.items():
                for pattern in patterns:
                    if pattern in header_text:
                        logger.debug(f"Identified table as '{table_type}' (found '{pattern}')")
                        return table_type
        except Exception as e:
            logger.debug(f"Error identifying table: {e}")
        
        return None
    
    def _fill_host_table(self, table, host: Host, host_index: int):
        """Fill host/inventory table."""
        # Prepare data rows
        data_rows = []
        
        # Basic info
        if host.hostname or host.ip:
            data_rows.append([
                "Имя хоста",
                host.hostname or "N/A",
                "IP адрес",
                host.ip or host.get_primary_ip() or "N/A"
            ])
        
        if host.os_name:
            os_full = host.os_name
            if host.os_version:
                os_full += f" {host.os_version}"
            if host.os_architecture:
                os_full += f" ({host.os_architecture})"
            data_rows.append(["Операционная система", os_full, "Роль в домене", host.domain_role or "N/A"])
        
        # Hardware
        if host.cpu or host.ram:
            hw_info = []
            if host.cpu:
                hw_info.append(f"CPU: {host.cpu}")
            if host.ram:
                hw_info.append(f"RAM: {host.ram}")
            if host.disk:
                hw_info.append(f"Disk: {host.disk}")
            data_rows.append(["Оборудование", " | ".join(hw_info), "", ""])
        
        # Clear existing data rows and add new ones
        self._clear_table_rows(table, keep_header=True)
        
        # Fill rows
        for i, row_data in enumerate(data_rows):
            if i < len(table.rows) - 1:
                row = table.rows[i + 1]
            else:
                row = table.add_row()
            for j, cell in enumerate(row.cells):
                if j < len(row_data):
                    cell.text = str(row_data[j]) if row_data[j] else ""
    
    def _fill_ports_table(self, table, port_scan: PortScan, host_index: int):
        """Fill open ports table."""
        ports = port_scan.open_ports
        
        if not ports:
            return
        
        # Clear existing data rows (keep header)
        self._clear_table_rows(table, keep_header=True)
        
        # Add rows for each port
        for port_data in ports:
            row = table.add_row()
            cells = row.cells
            
            # Map port data to columns based on header
            header = [c.text.lower() for c in table.rows[0].cells]
            
            for i, col in enumerate(header):
                if i >= len(cells):
                    break
                
                value = ""
                if 'port' in col or 'порт' in col:
                    value = port_data.get('port', '')
                elif 'protocol' in col or 'протокол' in col:
                    value = port_data.get('protocol', '')
                elif 'state' in col or 'состояние' in col or 'статус' in col:
                    value = port_data.get('state', '')
                elif 'service' in col or 'сервис' in col or 'служба' in col:
                    value = port_data.get('service', '')
                elif 'banner' in col or 'баннер' in col or 'версия' in col:
                    value = port_data.get('banner', '')
                
                cells[i].text = str(value)
        
        # Add SMB signing info if available
        if port_scan.smb_signing:
            row = table.add_row()
            row.cells[0].text = "SMB Signing"
            row.cells[1].text = port_scan.smb_signing
    
    def _fill_vulns_table(self, table, vulnerabilities: List[Vulnerability], host_index: int):
        """Fill vulnerabilities table."""
        if not vulnerabilities:
            return
        
        # Clear existing data rows
        self._clear_table_rows(table, keep_header=True)
        
        # Sort by severity
        severity_order = {'Critical': 0, 'High': 1, 'Medium': 2, 'Low': 3, 'Info': 4}
        sorted_vulns = sorted(
            vulnerabilities,
            key=lambda v: severity_order.get(v.severity, 5)
        )
        
        # Add rows for each vulnerability
        for vuln in sorted_vulns:
            row = table.add_row()
            cells = row.cells
            
            header = [c.text.lower() for c in table.rows[0].cells]
            
            for i, col in enumerate(header):
                if i >= len(cells):
                    break
                
                value = ""
                if 'cve' in col or 'идентификатор' in col or 'id' in col.lower():
                    value = vuln.cve or vuln.vuln_id
                elif 'severity' in col or 'критичность' in col or 'уровень' in col:
                    value = vuln.severity
                elif 'cvss' in col:
                    value = str(vuln.cvss_score) if vuln.cvss_score else ""
                elif 'title' in col or 'название' in col or 'описание' in col:
                    value = vuln.title or vuln.description
                elif 'exploit' in col or 'эксплойт' in col:
                    value = "Да" if vuln.exploit_available else "Нет"
                elif 'solution' in col or 'решение' in col or 'рекомендации' in col:
                    value = vuln.solution
                elif 'host' in col or 'хост' in col:
                    value = vuln.affected_host
                elif 'port' in col or 'порт' in col:
                    value = str(vuln.affected_port) if vuln.affected_port else ""
                elif 'service' in col or 'сервис' in col:
                    value = vuln.affected_service
                
                cells[i].text = str(value)[:500] if value else ""  # Truncate long values
    
    def _fill_software_table(self, table, software_list: List, host_index: int):
        """Fill installed software table."""
        self._clear_table_rows(table, keep_header=True)
        
        for sw in software_list:
            row = table.add_row()
            cells = row.cells
            header = [c.text.lower() for c in table.rows[0].cells]
            
            for i, col in enumerate(header):
                if i >= len(cells):
                    break
                
                if 'name' in col or 'название' in col or 'программа' in col:
                    cells[i].text = sw.name
                elif 'version' in col or 'версия' in col:
                    cells[i].text = sw.version
                elif 'date' in col or 'дата' in col:
                    cells[i].text = sw.install_date
    
    def _fill_services_table(self, table, services_list: List, host_index: int):
        """Fill services table."""
        self._clear_table_rows(table, keep_header=True)
        
        for svc in services_list:
            row = table.add_row()
            cells = row.cells
            header = [c.text.lower() for c in table.rows[0].cells]
            
            for i, col in enumerate(header):
                if i >= len(cells):
                    break
                
                if 'name' in col or 'название' in col or 'служба' in col:
                    cells[i].text = svc.name
                elif 'status' in col or 'состояние' in col or 'статус' in col:
                    cells[i].text = svc.status
                elif 'startup' in col or 'запуск' in col or 'тип' in col:
                    cells[i].text = svc.startup
                elif 'path' in col or 'путь' in col:
                    cells[i].text = svc.path
    
    def _fill_users_table(self, table, users_list: List, host_index: int):
        """Fill users/groups table."""
        self._clear_table_rows(table, keep_header=True)
        
        for ug in users_list:
            row = table.add_row()
            cells = row.cells
            header = [c.text.lower() for c in table.rows[0].cells]
            
            for i, col in enumerate(header):
                if i >= len(cells):
                    break
                
                if 'name' in col or 'имя' in col or 'пользователь' in col:
                    cells[i].text = ug.name
                elif 'type' in col or 'тип' in col:
                    cells[i].text = ug.type
                elif 'description' in col or 'описание' in col:
                    cells[i].text = ug.description
    
    def _fill_network_table(self, table, interfaces: List, host_index: int):
        """Fill network interfaces table."""
        self._clear_table_rows(table, keep_header=True)
        
        for iface in interfaces:
            row = table.add_row()
            cells = row.cells
            header = [c.text.lower() for c in table.rows[0].cells]
            
            for i, col in enumerate(header):
                if i >= len(cells):
                    break
                
                if 'name' in col or 'имя' in col or 'интерфейс' in col:
                    cells[i].text = iface.name
                elif 'ip' in col:
                    cells[i].text = iface.ip
                elif 'mask' in col or 'маска' in col:
                    cells[i].text = iface.mask
                elif 'gateway' in col or 'шлюз' in col:
                    cells[i].text = iface.gateway
                elif 'dns' in col:
                    cells[i].text = iface.dns
                elif 'mac' in col:
                    cells[i].text = iface.mac
    
    def _fill_updates_table(self, table, updates: List[str], host_index: int):
        """Fill updates/KB table."""
        self._clear_table_rows(table, keep_header=True)
        
        for kb in updates:
            row = table.add_row()
            if len(row.cells) > 0:
                row.cells[0].text = kb
    
    def _clear_table_rows(self, table, keep_header: bool = True):
        """Clear table rows, optionally keeping header."""
        start_idx = 1 if keep_header else 0
        
        # Collect rows to delete (can't modify while iterating)
        rows_to_delete = []
        for i in range(start_idx, len(table.rows)):
            rows_to_delete.append(i)
        
        # Delete in reverse order
        for i in reversed(rows_to_delete):
            try:
                tbl = table._tbl
                tr = tbl.trs[i]
                tr.getparent().remove(tr)
            except Exception as e:
                logger.debug(f"Error clearing row {i}: {e}")


def create_sample_template(output_path: Path):
    """Create a sample template document for testing."""
    doc = Document()
    
    # Title
    title = doc.add_heading('Протокол проверки информационной безопасности', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Intro text (should not be modified)
    doc.add_paragraph(
        "Данный документ содержит результаты автоматической проверки "
        "информационной безопасности инфраструктуры."
    )
    
    # Host inventory table
    doc.add_heading('1. Инвентаризация хостов', level=1)
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    headers = ['Параметр', 'Значение', 'Параметр', 'Значение']
    for i, h in enumerate(headers):
        header_cells[i].text = h
        header_cells[i].paragraphs[0].runs[0].bold = True
    
    # Ports table
    doc.add_heading('2. Открытые порты и сервисы', level=1)
    table = doc.add_table(rows=2, cols=5)
    table.style = 'Table Grid'
    headers = ['Порт', 'Протокол', 'Состояние', 'Сервис', 'Баннер/Версия']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    # Vulnerabilities table
    doc.add_heading('3. Уязвимости', level=1)
    table = doc.add_table(rows=2, cols=6)
    table.style = 'Table Grid'
    headers = ['CVE', 'Критичность', 'CVSS', 'Описание', 'Эксплойт', 'Рекомендации']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    # Software table
    doc.add_heading('4. Установленное ПО', level=1)
    table = doc.add_table(rows=2, cols=3)
    table.style = 'Table Grid'
    headers = ['Наименование', 'Версия', 'Дата установки']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    doc.save(str(output_path))
    logger.info(f"Sample template created: {output_path}")
